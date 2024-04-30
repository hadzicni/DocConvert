import os
from PyPDF2 import PdfFileReader
from docx import Document

def pdf_to_txt(pdf_file):
    with open(pdf_file, 'rb') as file:
        reader = PdfFileReader(file)
        text = ''
        for page_num in range(reader.numPages):
            page = reader.getPage(page_num)
            text += page.extractText()
    return text

def docx_to_txt(docx_file):
    doc = Document(docx_file)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def txt_to_pdf(txt_file, pdf_file):
    with open(txt_file, 'r', encoding='utf-8') as file:
        text = file.read()
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size = 12)
    pdf.multi_cell(200, 10, txt = text)
    pdf.output(pdf_file)

def txt_to_docx(txt_file, docx_file):
    document = Document()
    with open(txt_file, 'r', encoding='utf-8') as file:
        for line in file:
            document.add_paragraph(line.strip())
    document.save(docx_file)

def convert_file(input_file, output_file):
    file_extension = os.path.splitext(input_file)[1].lower()
    if file_extension == '.pdf':
        text = pdf_to_txt(input_file)
    elif file_extension == '.docx':
        text = docx_to_txt(input_file)
    elif file_extension == '.txt':
        with open(input_file, 'r', encoding='utf-8') as file:
            text = file.read()
    else:
        print("Nicht unterstütztes Dateiformat.")
        return

    output_extension = os.path.splitext(output_file)[1].lower()
    if output_extension == '.pdf':
        txt_to_pdf(text, output_file)
    elif output_extension == '.docx':
        txt_to_docx(text, output_file)
    elif output_extension == '.txt':
        with open(output_file, 'w', encoding='utf-8') as file:
            file.write(text)
    else:
        print("Nicht unterstütztes Ausgabeformat.")

if __name__ == "__main__":
    input_file = input("Geben Sie den Dateipfad der Eingabedatei ein: ")
    output_file = input("Geben Sie den Dateipfad der Ausgabedatei ein: ")
    convert_file(input_file, output_file)
